import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

style = document.styles['Heading 1']
font = style.font
font.name = 'Times'
font.size = Pt(36)

style = document.styles['Normal']
font = style.font
font.name = 'Times'
font.size = Pt(32)

volume_num = 1
head = 0
chapter_num_start = 1
chapter_num_end = 1
chapter_list = []
page = requests.get('https://www.wuxiaworld.com/novel/martial-god-asura')

#Create a BeautySoup object
soup = BeautifulSoup(page.text, 'html.parser')

volume_list = soup.find_all(class_="panel-body")

for volume in volume_list:
    
    if volume.find(class_="col-sm-6") == None:
        continue

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    document.add_picture('mga.jpg', width=Inches(8.27), height=Inches(11.69))
    document.add_section()

    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    chapter_name_list = volume.find_all(class_="chapter-item")
    #Get all the links from the div list
    chapter_links = []

    for chapter_link in chapter_name_list:
        chapter_links.append(chapter_link.find('a').get('href'))
        
    for chapter in chapter_links:
        link = 'https://www.wuxiaworld.com' + chapter
        page = requests.get(link)
        soup = BeautifulSoup(page.text, 'html.parser')
        story_view = soup.find_all(class_='p-15')
        for story_list in story_view:
            story_text = story_list.find_all('p')
            for story in story_text:
                chapter_list.append(story.get_text().replace('\xa0', ' ').replace('Previous Chapter', ''))
        for paragraph in chapter_list:
            if paragraph != '':
                if head == 0:
                    header = document.add_heading(paragraph)
                    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    head = 1
                else:
                    document.add_paragraph(paragraph)
        document.add_section()
        head = 0
        chapter_list = []
        print('Chapter ' + str(chapter_num_end) + ' down!')
        chapter_num_end+=1
        

    document.save('Martial God Asura Vol.' + str(volume_num) + ' ' + str(chapter_num_start) + '-' + str(chapter_num_end - 1) + '.docx')
    document = Document()

    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Times'
    font.size = Pt(36)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times'
    font.size = Pt(32)
    
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    print('Volume ' + str(volume_num) + ' down!')
    volume_num+=1
    chapter_num_start = chapter_num_end
