import os
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class HanaDocument(object):

    def __init__(self):
        self.document = Document()

    def addHead(self, text):
        header = self.document.add_heading(text)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def addPara(self, text):
        self.document.add_paragraph(text)

    def addSection(self):
        self.document.add_section()

    def sectionConfig(self, cm):
        for section in self.document.sections:
            section.top_margin = Cm(cm)
            section.bottom_margin = Cm(cm)
            section.left_margin = Cm(cm)
            section.right_margin = Cm(cm)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

    def stylesConfig(self, styleName, pt):
        style = self.document.styles[styleName]
        font = style.font
        font.name = 'Times'
        font.size = Pt(pt)

    def saveBook(self, name, volume, start, end):
        if not os.path.exists(name):
            os.mkdir(name)
            
        self.document.save(name + '/' + name + ' ' + str(volume) + ' ' + str(start) + '-' + str(end) + '.docx')
        
