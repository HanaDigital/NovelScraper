import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

document = Document()
sections = document.sections
head = 0
count = 1
page = requests.get('https://www.wuxiaworld.com/novel/martial-god-asura')

#Create a BeautySoup object
soup = BeautifulSoup(page.text, 'html.parser')

#Pull all text from the chapter-item div
chapter_name_list = soup.find_all(class_="chapter-item")

#Get all the links from the div list
chapter_links = []
for chapter_link in chapter_name_list:
    chapter_links.append(chapter_link.find('a').get('href'))

chapter_list = []

for chapter in chapter_links:
    link = 'https://www.wuxiaworld.com' + chapter
    page = requests.get(link)
    soup = BeautifulSoup(page.text, 'html.parser')
    story_view = soup.find_all(class_='p-15')
    for story_list in story_view:
        story_text = story_list.find_all('p')
        for story in story_text:
            chapter_list.append(story.get_text().replace('\xa0', ' ').replace('Previous Chapter', ''))
    for paragraph in chapter_list:
        if paragraph != '':
            if head == 0:
                document.add_paragraph(paragraph, style='TOCHeading')
                head = 1
            else:
                document.add_paragraph(paragraph)
    document.add_page_break()
    head = 0
    chapter_list = []
    print('Chapter ' + str(count) + 'down!')
    if(count == 3):
        break
    count+=1
    

document.save('MGAtest.docx')
    


